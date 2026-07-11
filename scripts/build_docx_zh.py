# -*- coding: utf-8 -*-
"""生成中文 Word 论文：RGF 体系详细介绍（含完整定理目录）。

用法：  python3 scripts/build_docx_zh.py
依赖：  python-docx
输出：  论文_RGF体系详解.docx （项目根目录）
"""
import os, json
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.join(os.path.dirname(__file__), '..')
decls = json.load(open(os.path.join(ROOT, 'scripts', 'decls.json'), encoding='utf-8'))
headers = json.load(open(os.path.join(ROOT, 'scripts', 'headers.json'), encoding='utf-8'))

doc = Document()

# ---------- 基础字体：中西文分设 ----------
HANS = '宋体'         # 正文中文
HANS_HEI = '黑体'     # 标题中文
LATIN = 'Times New Roman'
MONO = 'Consolas'


def set_style_cjk(style, latin=LATIN, cjk=HANS, size=None):
    style.font.name = latin
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), latin)
    rfonts.set(qn('w:hAnsi'), latin)
    rfonts.set(qn('w:eastAsia'), cjk)
    if size:
        style.font.size = Pt(size)


set_style_cjk(doc.styles['Normal'], LATIN, HANS, 11)
for s in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Title']:
    try:
        set_style_cjk(doc.styles[s], LATIN, HANS_HEI)
    except KeyError:
        pass


def _run_cjk(run, latin=LATIN, cjk=HANS):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), latin)
    rfonts.set(qn('w:hAnsi'), latin)
    rfonts.set(qn('w:eastAsia'), cjk)


def add_body(text, size=11, italic=False, color=None, before=2, after=4, cjk=HANS, first_indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(22)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)
    _run_cjk(r, LATIN, cjk)
    return p


def add_code(text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = MONO
    r.font.size = Pt(size)
    _run_cjk(r, MONO, MONO)
    return p


def H(text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        _run_cjk(r, LATIN, HANS_HEI)
    return h


def add_theorem(entry):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    kind = '定理' if entry['kind'] == 'theorem' else '引理'
    rk = p.add_run(f'[{kind}] ')
    rk.font.size = Pt(9); rk.bold = True; rk.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
    _run_cjk(rk, LATIN, HANS_HEI)
    rn = p.add_run(entry['name'])
    rn.font.name = MONO; rn.font.size = Pt(9); rn.bold = True
    _run_cjk(rn, MONO, MONO)
    add_code(entry['sig'])
    if entry.get('doc'):
        pd = doc.add_paragraph()
        pd.paragraph_format.space_before = Pt(0); pd.paragraph_format.space_after = Pt(2)
        pd.paragraph_format.left_indent = Inches(0.25)
        rd = pd.add_run('原文档字符串（source docstring）: ' + entry['doc'][:500])
        rd.italic = True; rd.font.size = Pt(8); rd.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        _run_cjk(rd, LATIN, HANS)


# 统计量
total = sum(len(v) for v in decls.values())
tcount = sum(1 for v in decls.values() for x in v if x['kind'] == 'theorem')
lcount = total - tcount
nfiles = len(decls)

# ================= 封面 =================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('递归生成框架（RGF）及其内部超穷集合论 RGF 2.0')
r.bold = True; r.font.size = Pt(22); _run_cjk(r, LATIN, HANS_HEI)
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = st.add_run('体系详解与完整定理目录'); rr.bold = True; rr.font.size = Pt(16); _run_cjk(rr, LATIN, HANS_HEI)
s3 = doc.add_paragraph(); s3.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr3 = s3.add_run('——一个基于 Lean 4 + Mathlib 的机器可验证形式化体系'); rr3.font.size = Pt(12); _run_cjk(rr3, LATIN, HANS)
doc.add_paragraph()
s4 = doc.add_paragraph(); s4.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr4 = s4.add_run(f'（全库 {nfiles} 个 .lean 源文件，共 {total} 条定理/引理；机器整体验证通过）')
rr4.font.size = Pt(10); rr4.italic = True; _run_cjk(rr4, LATIN, HANS)

# ---------- 摘要 ----------
doc.add_paragraph()
H('摘要', 1)
add_body(
    '本文系统、详细地介绍递归生成框架（Recursive Generative Framework，RGF）及其内部超穷集合论升级版 '
    'RGF 2.0，并给出整个 Lean 4 形式化库中每一条定理的完整目录。RGF 以“递归生成”为唯一原语：任何结构都被'
    '视为沿某条时间流反复迭代同一个动作 step = modify ∘ generate（先“生成”一个候选对象，再依据某个不变量对其'
    '“修正”）的产物。自此原语出发，本体系在 Lean 4 + Mathlib 之上依次构造自然数、整数、有理数与实数，并进一步'
    '展开出涵盖分析、代数、范畴论、谱论、图论与拓扑的数学脊柱；在物理侧，同一底层通过递归本构动力学（Recursive '
    'Constitutive Dynamics，RCD）展开为时空、引力、规范结构与标准模型参数谱的涌现。',
    first_indent=True)
add_body(
    'RGF 2.0 是针对“如何把 RGF 的内部集合论从有限体系（ZF−∞，强度约等于皮亚诺算术 PA）无缝升级到完整 ZFC”'
    '这一问题给出的四条技术路径的形式化实现：升维编码域（W-类型树商）、带反射原理的内部累积层级、带强迫骨架的'
    '布尔值模型、以及超穷生成流（超穷递归、Hartogs 函数、良序化与选择公理）。所有结论均为真正的机器证明：全库'
    '通过 lake build，无 sorry（源码中仅少量有限 Coxeter 群阶缺口被诚实标注为 OPEN），无自定义 axiom，无 '
    '@[implemented_by]；头部命题经 #print axioms 审计，其内核公理足迹完全落在标准允许集合之内。',
    first_indent=True)
add_body('关键词：递归生成框架；递归本构动力学；形式化数学；Lean 4；内部集合论；ZFC；W-类型；布尔值模型；'
         '强迫法；超穷递归；涌现；机器验证。', italic=True)

# ================= 第一部分 =================
doc.add_page_break()
H('第一部分  体系综述', 1)

H('1. 引言：一个原语，一个系统', 2)
add_body(
    'RGF 的出发点是一个哲学上极简、数学上可完全形式化的立场：任何结构都是单一动作“递归生成”的产物。'
    '生成动作记作 step = modify ∘ generate——先“生成”一个新的候选对象，再依据某个不变量对其“修正”。'
    '自然数是该动作产出的第一批对象（zero | succ 归纳）；整数、有理数与实数随后经由标准代数构造逐级产生；'
    '再往上，分析、代数、范畴论、谱论、图论与拓扑都是同一生成动作在不同层次上的产物。', first_indent=True)
add_body(
    '在物理侧，RGF 的动力学实现被称为递归本构动力学（RCD）：沿时间流迭代离散的“锁膜/恢复”规则，在粗粒化与'
    '重整化流之下，产生连续时空、引力、规范结构与标准模型参数谱的涌现。RGF 与 RCD 共享同一套经过验证的数学'
    '底层，因此“数学基础”与“物理唯象”在同一个形式系统内部闭合。', first_indent=True)

H('2. 系统架构与四层结构', 2)
add_body('整个体系是一个名为 RGF 的 Lean 库，按主题组织为四个层级，目录结构与命名空间严格对应：', first_indent=True)
add_body('· 第一层，生成脊柱（RGF.Generative.*）：生成原语本身、锁定、唯一性与元理论。', before=0, after=0)
add_body('· 第二层，数学产物（RGF.Math.*）：内部数系、集合论、分析、范畴论、谱论、图论、代数、拓扑——'
         '皆为第一层生成动作的产物。', before=0, after=0)
add_body('· 第三层，物理动力学（RGF.Physics.*）：普适物理动力学与涌现（RCD 的形式化）。', before=0, after=0)
add_body('· 第四层，唯象学（RGF.Phenomenology.*）：标准模型唯象学与测试/审计套件。', before=0, after=0)
add_body('此外还有一个应用层（RGF.Applications.*），把生成动力学应用到遍历性、轨道计数、内蕴对称性等具体问题。')
add_body(
    '各层维持严格的单向依赖：第一层的中性生成核心（RGF.Generative.Core.*）连同数学底层（RGF.Math.*）构成'
    '体系的 Tier 1，其传递 import 闭包中不含任何物理/唯象模块。这种“Tier 1 的单向独立性”由脚本'
    '（check_tier_boundary.py）机械检查，并配有反例自测保证该关卡有效，因此“数学基础不依赖物理假设”不是口号，'
    '而是一个可验证的架构不变量。', first_indent=True)

H('3. 生成脊柱与内部数系', 2)
add_body(
    '生成原语首先构造自然数 RGFNat（zero | succ 归纳），再经标准代数构造逐级产出整数 RGFInt、有理数 RGFRat 与'
    '实数 RGFReal′。RGFReal′ 由 RGF 有理数的柯西序列构造，并被证明与标准实数完全一致：', first_indent=True)
add_code('noncomputable def orderedRingEquivReal : RGFReal′ ≃+*o ℝ', size=10)
add_body(
    '这条有序环同构（RGF.RGFReal′.orderedRingEquivReal）表明：RGF 内部生成的实数在结构上与经典实数不可区分，'
    '因此由生成原语构建的分析对象能与 Mathlib 的实分析无缝对接。这使 RGF 的“自建数系”既是一个本体论上的生成'
    '产物，又在数学上等价于标准实数——这是整个体系的第一块基石。', first_indent=True)

H('4. 内部超穷集合论：从生成树到完整 ZFC（RGF 2.0 的四条路径）', 2)
add_body(
    'RGF 1.0 的内部集合论使用阿克曼编码把集合映射为有限二进制整数，因而只能容纳遗传有限集（HF），一致性强度'
    '约等于皮亚诺算术（PA）。RGF 2.0 通过对“编码域、生成原语、逻辑值域”的根本性重构，把内部集合论从有限体系'
    '无缝升级到完整 ZFC。以下四条路径与体系中的模块一一对应。', first_indent=True)

H('路径一  升维编码域：从算术编码到无损 W-类型', 3)
add_body(
    '集合论核心的编码域采用带标号的良基归纳树（W-类型）：一个集合就是一棵分枝树，其“外延相等”由树之间的双向'
    '嵌入给出，属于关系 Mem₂ 定义在树的分枝上。对这些树按等价关系取商，得到内部集合宇宙 RGFSet₂。此时实数不再'
    '是集合编码的载体，而成为树状集合宇宙中的一个普通内部对象。在该宇宙上，体系一举证明完整的十条 ZFC 公理：',
    first_indent=True)
add_code('theorem RGF2_models_ZFC :\n'
         '  (extensionality) ∧ (empty set) ∧ (pairing) ∧ (union) ∧ (power set) ∧\n'
         '  (separation schema) ∧ (replacement schema) ∧ (foundation) ∧ (infinity) ∧ (choice)', size=9)
add_body('对应模块：RGF/Math/SetTheory/RGF2/Core/WType.lean、RGF2.lean、Master.lean。定理 '
         'infinity_strict_upgrade 进一步表明 RGF 2.0 满足无穷公理，而 RGF 1.0 的遗传有限核心可证满足其否定，'
         '确认这是一次严格升级。')

H('路径二  内部宇宙层次与反射原理', 3)
add_body(
    '为克服哥德尔不完备性带来的“低一致性强度”局限，体系在 RGF 内部构造累积层级 Vhier：V₀=∅，V(α+1)=P(Vα)，'
    '极限层取并 Vλ=⋃_{β<λ}Vβ。围绕它形式化了穷尽性、极限层并集方程、秩演算以及 Montague–Lévy 反射原理，'
    '并利用 Lean 4 的多级宇宙机制表达 Vω 乃至不可达基数宇宙 Vκ。定理 RGF2_inaccessible_gives_ZFC_model 诚实地'
    '把结论写成条件形式：若存在强不可达基数，则存在 ZFC 的传递模型——把“ZFC + 不可达”置于 ZFC 之上，而非声称'
    ' RGF 本身超越 ZFC。', first_indent=True)
add_body('对应模块：RGF/Math/SetTheory/RGF2/Hierarchy/（Cumulative、Reflection、Inaccessible、Transfinite、'
         'HereditarilyFinite 等）。')

H('路径三  拓宽逻辑值域：从二值逻辑到布尔值模型', 3)
add_body(
    '为消除模型刚性、获得强迫法的弹性，体系把属于关系的值域从二值 Prop 泛化为任意完备布尔代数 B，构造布尔值'
    '宇宙 BSet B（即 V^B）。定理 RGF2_boolean_valued_ZFC_full 证明每条 ZFC 公理在布尔值 ⊤ 处成立；'
    'RGF2_forcing_relation_calculus 给出强迫关系 ⊩ᴮ 的标准演算律；IndependenceFamily(Infinite) 给出独立性命题'
    '（如 ¬CH）的组合骨架。由此，集合论的独立性命题可直接在布尔值模型内部讨论。', first_indent=True)
add_body('对应模块：RGF/Math/SetTheory/RGF2/Boolean/（Model、ValuedZFC、Forcing、ForcingRelation、'
         'ChoiceReplacement、IndependenceFamily、IndependenceFamilyInfinite）。')

H('路径四  在生成算子中内化超穷递归与良序原理', 3)
add_body(
    'RGF 当前的“生成”原语沿自然数或实数时间轴推进，在认识论上被限制在 ℵ₀ 或 2^ℵ₀ 步之内。路径四把迭代算子'
    '推广到超穷时间流，使生成动作能沿任意良序集作超穷归纳推进；并在 RGF 树上形式化 Hartogs 定理——对任意集合 X '
    '生成一个比 X 更大的序数——以此为基础在 RGF 内部证明良序化定理与（超穷）选择公理。诸如 '
    'tfIterate_omega_eq_lim 之类的定理进一步证明极限阶规则确实是“承重”的（改变极限规则即改变结论），'
    '打消了“极限层只是平凡改名”的疑虑。', first_indent=True)
add_body('对应模块：RGF/Math/SetTheory/RGF2/Hierarchy/Transfinite.lean、TransfiniteLimitLoadBearing.lean、'
         'Core/InternalOrdinal.lean、Core/InternalCardinal.lean。')

H('5. 锁定、唯一性与诚实性', 2)
add_body(
    '生成脊柱的一个核心议题是“相位锁定”：递归系统在何种条件下锁定到特定的对称性与维数。体系形式化了 Z5 相位'
    '锁定机制，并对“为什么是 5”这一关键问题给出诚实、不夸大的结论：', first_indent=True)
add_code('theorem k_five_conditional_on_solvability_criterion :\n'
         '  (∀ sys, FiveLockingCondition sys → sys.k = 5) ∧\n'
         '  (∀ sys, EmergenceCondition sys ↔ ¬ IsSolvable (Equiv.Perm (Fin sys.k)))', size=9)
add_body(
    '该定理诚实地指出：一旦给定“可解性判据”，k=5 便由极小性唯一地被迫出；但该判据本身是以定义方式引入的。'
    '类似地，空间的三维锁定来自代数自对偶 dim 𝔰𝔬(d)=d（即 d(d−1)=2d），对 d≥1 而言当且仅当 d=3 成立。'
    '唯一性层（RGF.Generative.Uniqueness.*）系统地区分“可证”与“不可证”，把建模假设与推导结论分离开来。',
    first_indent=True)

H('6. 递归本构动力学（RCD）与物理涌现', 2)
add_body('在共享的数学底层之上，RCD 沿时间流迭代离散的锁膜规则，并考察其在粗粒化/重整化流下的有效涌现。'
         '体系为若干物理命题给出机器验证的核心，例如：', first_indent=True)
add_body('· 空间维数锁定 spatial_dimension_three；量子关联上界 chsh_tsirelson_bound（Tsirelson 界 2√2）。',
         before=0, after=0)
add_body('· 从格点到连续：由锁膜格点规则导出时域递归方程，含扩散系数、粗粒化噪声与 FORS 核。', before=0, after=0)
add_body('· 普朗克尺度 planckLength_ratio_bounds；带宇宙学涌现与暗能量系数的递归/螺旋标度律。', before=0, after=0)
add_body('· 弦与引力子、解锁模相变，以及时域粗粒化的信息定域化相图。', before=0, after=0)

H('7. 标准模型唯象学', 2)
add_body('唯象层把上述结构推进到标准模型参数谱。规范群涌现给出一个可验证的定量结论——大统一尺度处的 Weinberg 角：',
         first_indent=True)
add_code('theorem weinberg_angle_gut : (∑ i, weakIsospin i ^2)/(∑ i, charge i ^2) = 3/8', size=10)
add_body('该定理证明弱同位旋平方和与电荷平方和之比恰为 3/8，即经典的 sin²θ_W = 3/8。本层还包含超荷无迹性、'
         '五极相位和、中微子混合（如三双极大 θ₁₃ 的证伪）等命题，并配有测试/审计套件。', first_indent=True)

H('8. 验证方法论与可信性', 2)
add_body('体系把“可信性”本身作为一等目标，采用如下机制：', first_indent=True)
add_body('· 无 sorry、无自定义 axiom、无 @[implemented_by]：全库通过 lake build，所有结论均为真正的机器证明'
         '（唯一被诚实标注为 OPEN 的缺口是有限 Coxeter 群 |W_H3|、|W_H4| 的群阶，Mathlib 尚不含相应分类定理）。',
         before=0, after=0)
add_body('· 可审计的公理足迹：头部命题经 #print axioms 验证，只依赖 propext、Classical.choice、Quot.sound'
         '（少数有限判定另用 Lean.ofReduceBool、Lean.trustCompiler）。', before=0, after=0)
add_body('· 机械化的架构不变量：Tier 1 的单向独立性由脚本检查，并配反例自测保证关卡有效。', before=0, after=0)
add_body('· 诚实性定理：尚未闭合的建模假设以条件定理与 Iff.rfl 形式明确标注，不把定义冒充为推导。',
         before=0, after=0)

H('9. 规模统计（自源码树运行时提取）', 2)
add_body(f'· 源文件总数：{nfiles} 个 .lean 文件。', before=0, after=0)
add_body(f'· 定理/引理总数：{total} 条（其中定理 {tcount} 条，引理 {lcount} 条）。', before=0, after=0)
add_body('· 组织为 4 个主题层级 + 1 个应用层，共若干子目录（逐目录、逐文件的目录见第二部分）。', before=0, after=0)
add_body('说明：下方第二部分逐文件列出每一条定理/引理的名称与 Lean 类型签名（截断到证明起始 := 之前），'
         '若源码带有文档字符串则附上原始描述，以便读者在源码中逐条检索与交叉核对。', italic=True)

# ============ Tier / Dir 中文导言 ============
TIER_INTRO = {
 'RGF/Generative': '生成脊柱是体系的第一层，形式化“递归生成”原语本身及其后果：核心语法与不变量、相位/维数锁定'
    '机制、初始假设的必要性与唯一性，以及关于生成谓词之可证伪性、普适性等的元理论。',
 'RGF/Math': '数学产物层是第一层生成动作的直接产物：内部数系与实数、内部超穷集合论 RGF 2.0、分析、代数、'
    '范畴论、谱论、图论与拓扑。它连同生成核心构成不依赖物理假设的 Tier 1。',
 'RGF/Physics': '物理动力学层是递归本构动力学（RCD）的形式化：粗粒化、重整化群、标度律、格点到连续的极限、'
    '时空与引力涌现、普朗克尺度与宇宙学等。',
 'RGF/Phenomenology': '唯象层把生成动力学推进到标准模型参数谱：规范群涌现、Weinberg 角、中微子混合、Koide '
    '关系等，并配有测试与审计套件。',
 'RGF/Applications': '应用层把生成动力学应用到具体数学问题：构造性遍历性、轨道计数、线性化谱、内蕴对称性、'
    'Babai 型弱结果等。',
}

DIR_INTRO = {
 'RGF/Generative/Assembly': '装配层：把核心命题（六大命题、命题 C 等）装配为整体结论与简化核心。',
 'RGF/Generative/Bridge': '桥接层：把抽象 RGF 动力学与标准 Mathlib 结构连接，含两层/三层普适性、对称性推导、'
    '可移植 RGF 态与动力学性质。',
 'RGF/Generative/Core': '生成核心：递归语法、基本定义与不变量、求值泛函、RGF 公理与核心命题（Paper1 RCE 基础、'
    '自指孤子等）。这是整个体系的中性底层。',
 'RGF/Generative/Locking': '锁定层：相位锁定与维数锁定的完整理论——Z5 相位锁定、五重对称、临界比、三维唯一性、'
    '模分解与模锁定、五重锁定唯一性、Penrose 准晶/K-理论等。',
 'RGF/Generative/Meta': '元理论层：生成谓词的可证伪性、RGF 普适性、算法信息、解析数论、Lomb–Scargle 与数值验证、'
    'RGF 与黎曼猜想之关系的现状等。',
 'RGF/Generative/Uniqueness': '唯一性层：初始假设的必要性与极小性、复基域的必要性、覆盖性、G1 排他性、'
    '可实现性障碍、稳定性唯一性——系统地区分“可证”与“不可证”。',
 'RGF/Math/Algebra': '代数：单群 A5、仿射 Kac–Moody、代数几何、Coxeter 五重对称、晶体学限制、二面体群、'
    '有限群傅里叶分析、李代数嵌入、非交换几何、PSL(2,7)、Clifford 手征、模形式与表示论等。',
 'RGF/Math/Analysis': '分析：构造性 Feynman–Kac 与常微分方程、连续极限、收敛理论、微积分基本定理、泛函分析、'
    'Γ-收敛、归纳极限、测度与积分、RGF 连续性/导数/极限、严格拟凸性等。',
 'RGF/Math/Category': '范畴论：动力学 HoTT 及其流、等变函子、不动点子范畴、遗忘函子、生成动力学范畴与生成精化、'
    'RGF 范畴。',
 'RGF/Math/Graph': '图论：Frucht 定理（一般/RGF 版）、图自同构、轨道分类与配对、Petersen 图、量子轨道配对、'
    'Burnside、着色、平面性、Ramsey、t-设计、Turán 图。',
 'RGF/Math/Real': '内部数系与实数：RGFNat/Int/Rat 及其序与域结构、柯西序列与 Dedekind 分割、Banach 压缩、'
    'RGFReal′ 与标准实数的有序环同构、内部无穷、积分等。这是体系的第一块基石。',
 'RGF/Math/SetTheory': '集合论主目录：RGF–ZFC 对比、一致性强度、Galois 连接、Grothendieck 宇宙、ZF 基线，'
    '以及 RGF 2.0 子系统的入口。',
 'RGF/Math/SetTheory/RGF2': 'RGF 2.0 主体：W-类型树商宇宙 RGFSet₂、十条 ZFC 公理的主定理 Master、涌现的 ω、'
    '生成桥接，以及相对强度的诚实标注。',
 'RGF/Math/SetTheory/RGF2/Boolean': '布尔值模型（路径三）：完备布尔代数取值的宇宙 V^B、布尔值 ZFC、'
    '强迫关系及其演算、选择与替换，以及 ¬CH 独立性族的组合骨架。',
 'RGF/Math/SetTheory/RGF2/Core': 'RGF 2.0 内核：W-类型定义、内部序数与内部基数、对称差群结构。',
 'RGF/Math/SetTheory/RGF2/Hierarchy': '累积层级与超穷（路径二与路径四）：Vhier 累积层级、反射原理、不可达基数'
    '与 ZFC 模型、超穷递归与承重的极限阶、遗传有限层。',
 'RGF/Math/Spectral': '谱论：Cheeger 不等式、热核、拉普拉斯算子、Rayleigh 商、谱隙动力学、谱方法、'
    '标准模型谱三元组、谱约化猜想。',
 'RGF/Math/Topology': '拓扑：代数拓扑、边界路径、配边假设、构造性 TDA、de Rham、微分几何、Euler 示性数、'
    '填充半径、Fisher 度量、持续同调、Regge 演算、拓扑涌现。',
 'RGF/Physics/Dynamics': '动力学：粗粒化与重整化群、反常标度、双层对偶、共形场论、细致平衡、离散规范理论与迹公式、'
    '容错阈值、流体动力学极限、KPZ 及其推导、混合时间、相图、正则性结构、随机 PDE 的 FRG、环面码、'
    'Yau 相对熵等。',
 'RGF/Physics/Emergence': '涌现：完整推导链与唯一性、复杂度/能量/熵涌现、FORS 能量第一性原理、几何约化、'
    '格点到连续与色散、时空与引力（Paper2）、普朗克尺度（Paper7）、宇宙学（Paper9）、弦与引力子、'
    '递归生成的常数等。',
 'RGF/Phenomenology/StandardModel': '标准模型：α/β 参数、失谐分析、大统一、Koide 扩展、非厄米轻子相位、'
    '中微子预言与 σ 偏差、规范涌现与分配（Paper6）、普朗克唯象、规范比、生成常数、阈值、ζ 正规化、'
    '26 参数谱等。',
 'RGF/Phenomenology/TestSuite': '测试与审计套件：完整审计、去重审计、native_decide 审计、KAM 丢番图判据、'
    '原创数学与前沿命题、对批评的回应，以及各批次新增/追加的定理集合。',
}

# ================= 第二部分：完整目录 =================
doc.add_page_break()
H('第二部分  逐模块细节与完整定理目录', 1)
add_body('本部分按四个主题层级 + 应用层的顺序，逐目录、逐文件列出源码中的全部定理与引理。每一条给出其种类、'
         '名称与 Lean 类型签名；若源码带有文档字符串，则附上原始描述（截断显示）。', first_indent=True)

tiers = [
    ('RGF/Generative', 'Tier 1  生成脊柱 RGF.Generative'),
    ('RGF/Math', 'Tier 2  数学产物 RGF.Math'),
    ('RGF/Physics', 'Tier 3  物理动力学 RGF.Physics'),
    ('RGF/Phenomenology', 'Tier 4  标准模型唯象学 RGF.Phenomenology'),
    ('RGF/Applications', '应用层  RGF.Applications'),
]

allfiles = sorted(decls.keys())

for tier_prefix, tier_title in tiers:
    H(tier_title, 2)
    if tier_prefix in TIER_INTRO:
        add_body(TIER_INTRO[tier_prefix])
    tfiles = [f for f in allfiles if f.startswith(tier_prefix + '/')]
    bydir = OrderedDict()
    for f in tfiles:
        d = os.path.dirname(f)
        bydir.setdefault(d, []).append(f)
    for d in sorted(bydir):
        H(d.replace('RGF/', ''), 3)
        if d in DIR_INTRO:
            add_body(DIR_INTRO[d])
        for f in bydir[d]:
            fname = os.path.basename(f)
            entries = decls[f]
            H(fname + f'  （{len(entries)} 条）', 4)
            hdr = headers.get(f, '')
            if hdr:
                add_body('模块说明（source header）: ' + hdr[:400], size=9, italic=True,
                         color=(0x33, 0x33, 0x66), before=0, after=2)
            for e in entries:
                add_theorem(e)

out = os.path.join(ROOT, '论文_RGF体系详解.docx')
doc.save(out)
print('saved', out, '| files', nfiles, '| decls', total)
