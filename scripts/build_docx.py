# -*- coding: utf-8 -*-
"""Build the detailed Chinese Word document with the full theorem catalog."""
import os, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.join(os.path.dirname(__file__), '..')
decls = json.load(open(os.path.join(ROOT, 'scripts', 'decls.json'), encoding='utf-8'))
headers = json.load(open(os.path.join(ROOT, 'scripts', 'headers.json'), encoding='utf-8'))

doc = Document()

# ---- base styles: set an East-Asian friendly font ----
def set_cjk(style, font='Times New Roman', size=None):
    style.font.name = 'Times New Roman'
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font)
    if size:
        style.font.size = Pt(size)

set_cjk(doc.styles['Normal'], 'Times New Roman', 11)
for s in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Title']:
    try:
        set_cjk(doc.styles[s], 'Arial')
    except KeyError:
        pass

MONO = 'Consolas'

def add_body(text, size=11, italic=False, color=None, before=2, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

def add_code(text, size=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = MONO
    r.font.size = Pt(size)
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), MONO); rfonts.set(qn('w:hAnsi'), MONO)
    return p

def add_theorem(entry):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    kind = 'Theorem' if entry['kind'] == 'theorem' else 'Lemma'
    rk = p.add_run(f'[{kind}] ')
    rk.font.size = Pt(9); rk.bold = True; rk.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
    rn = p.add_run(entry['name'])
    rn.font.name = MONO; rn.font.size = Pt(9); rn.bold = True
    rpr = rn._element.get_or_add_rPr(); rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'), MONO); rf.set(qn('w:hAnsi'), MONO); rpr.append(rf)
    add_code(entry['sig'])
    if entry.get('doc'):
        pd = doc.add_paragraph()
        pd.paragraph_format.space_before = Pt(0); pd.paragraph_format.space_after = Pt(2)
        pd.paragraph_format.left_indent = Inches(0.25)
        rd = pd.add_run('Note (source docstring): ' + entry['doc'][:500])
        rd.italic = True; rd.font.size = Pt(8); rd.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ================= TITLE =================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('The Recursive Generative Framework (RGF) and its Internal Transfinite Set Theory RGF 2.0')
r.bold = True; r.font.size = Pt(20); set_run_cjk = None
r2 = doc.add_paragraph(); r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = r2.add_run('System Overview and Full Theorem Catalogue'); rr.bold = True; rr.font.size = Pt(16)
r3 = doc.add_paragraph(); r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr3 = r3.add_run('A complete machine-verified formalization based on Lean 4 + Mathlib'); rr3.font.size = Pt(12)
doc.add_paragraph()

total = sum(len(v) for v in decls.values())
tcount = sum(1 for v in decls.values() for x in v if x['kind']=='theorem')
lcount = total - tcount

# placeholder; real prose injected below via SECTIONS
exec(open(os.path.join(ROOT, 'scripts', 'prose.py'), encoding='utf-8').read())

# ================= PART II: full catalog =================
doc.add_page_break()
h = doc.add_heading('Part II  Module-by-module details and full theorem catalogue', level=1)

# ordering of tiers and dirs
import re
def natural_files():
    order = []
    for dirpath, dirs, files in os.walk(os.path.join(ROOT, 'RGF')):
        pass
    return order

# Build mapping dir -> files (sorted), preserving tier grouping
tiers = [
    ('RGF/Generative', 'Tier 1  Generative backbone RGF.Generative'),
    ('RGF/Math', 'Tier 2  Mathematical products RGF.Math'),
    ('RGF/Physics', 'Tier 3  Physical dynamics RGF.Physics'),
    ('RGF/Phenomenology', 'Tier 4  Standard Model phenomenology RGF.Phenomenology'),
    ('RGF/Applications', 'Applications tier  RGF.Applications'),
]

allfiles = sorted(decls.keys())

for tier_prefix, tier_title in tiers:
    doc.add_heading(tier_title, level=2)
    if tier_prefix in TIER_INTRO:
        add_body(TIER_INTRO[tier_prefix])
    # group by immediate subdirectory
    tfiles = [f for f in allfiles if f.startswith(tier_prefix + '/')]
    # subdir key = path minus filename
    from collections import OrderedDict
    bydir = OrderedDict()
    for f in tfiles:
        d = os.path.dirname(f)
        bydir.setdefault(d, []).append(f)
    for d in sorted(bydir):
        doc.add_heading(d.replace('RGF/', ''), level=3)
        if d in DIR_INTRO:
            add_body(DIR_INTRO[d])
        for f in bydir[d]:
            fname = os.path.basename(f)
            entries = decls[f]
            ph = doc.add_heading(fname + f'  ({len(entries)} entries)', level=4)
            hdr = headers.get(f, '')
            if hdr:
                add_body('Module description: ' + hdr[:400], size=9, italic=True, color=(0x33,0x33,0x66), before=0, after=2)
            for e in entries:
                add_theorem(e)

out = os.path.join(ROOT, 'RGF_System_Overview_and_Full_Theorem_Catalog.docx')
doc.save(out)
print('saved', out, 'total decls', total)
